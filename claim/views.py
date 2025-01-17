from http.client import HTTPResponse
from django.shortcuts import render
from claim.models import Claim
from pathlib import Path
import os
from django.http import HttpResponse
import random
import docx
import xlsxwriter
import openpyxl

# Create your views here.
BASE_DIR = Path(__file__).resolve().parent.parent

def create_report(claims):
    SAVE_NAME = str(random.randint(1, 1001)) + '.docx'

    try:
        dir = os.path.abspath(str(BASE_DIR) + '/static/template_doc/программа_template.docx')
        currentDocument = docx.Document(dir)
    except FileNotFoundError:
        print('FileNotFoundError...')
        return
    
    table = currentDocument.add_table(rows=len(claims), cols=3)
    table.style = 'poigraem_table'

    for row in table.rows:
        row.cells[0].width = docx.shared.Inches(0.3)
        row.cells[1].width = docx.shared.Inches(3.5)
    
    for i, claim in enumerate(claims):
        hdr_cells = table.rows[i].cells
        hdr_cells[0].paragraphs[0].add_run(str(i + 1)).bold = False
        hdr_cells[1].paragraphs[0].add_run(claim[0]).bold = True
        hdr_cells[1].paragraphs[0].add_run('\n' + claim[1] + ', ')
        hdr_cells[1].paragraphs[0].add_run(str(claim[2]) + ' кл. ')
        hdr_cells[1].paragraphs[0].add_run('(' + str(claim[3]) + ')')
        hdr_cells[1].paragraphs[0].add_run('\nпреп. ' + claim[4])
        if claim[5] != None:
            hdr_cells[1].paragraphs[0].add_run('\nконц. ' + claim[5])
        hdr_cells[2].text = claim[6]
    try:
        dir = os.path.abspath(str(BASE_DIR) + ('/static/temp_files/' + SAVE_NAME))
        currentDocument.save(dir)
        dirs = [SAVE_NAME, '/static/temp_files/' + SAVE_NAME]
        return dirs
    except:
        print('Dont save file...')
        return 'Error'

def index(request):
    if request.method == 'POST':
        print('POST POST POST')
        type_cl = request.POST.get('type_claim')
        select_cl = request.POST.get('select_claim')
        name_cl = request.POST.get('name_claim')
        parent_cl = request.POST.get('parent_claim')
        class_cl = request.POST.get('class_claim')
        concert_cl = request.POST.get('concert_claim')
        city_cl = request.POST.get('city_claim')
        prog_cl = request.POST.get('prog_claim')
        url_cl = request.POST.get('url_claim')

        data = {'type': type_cl, 'select': select_cl, 'name': name_cl,
                'parent': parent_cl, 'class': class_cl, 'concert': concert_cl,
                'city': city_cl, 'prog': prog_cl, 'url': url_cl
                }
        claim = Claim.objects.create(type_claim=type_cl,
                                    select_claim=select_cl,
                                    name_claim=name_cl,
                                    parent_claim=parent_cl,
                                    class_claim=class_cl,
                                    concert_claim=concert_cl,
                                    city_claim=city_cl,
                                    prog_claim=prog_cl,
                                    url_claim=url_cl,
                                    place_claim=0)
        data['id'] = claim.id
        return render(request, 'apply.html', context=data)
    return render(request, 'index.html')

def admin(request):
    claims = Claim.objects.all()
    return render(request, 'admin.html', context={'claims': claims})

def generate(request):
    if request.method == 'POST':
        claims = []
        xl_file = openpyxl.load_workbook(request.FILES['fileXlsx'])
        sheet = xl_file['Data']
        for i, row in enumerate(sheet.rows):
            if i == 0:
                continue
            claim = []
            claim.append(row[7].value)
            claim.append(row[3].value)
            claim.append(row[5].value)
            claim.append(row[2].value)
            claim.append(row[4].value)
            claim.append(row[6].value)
            claim.append(row[8].value)
            claims.append(claim)
        dirs = create_report(claims)
        return render(request, 'word.html', context={'dir': dirs})
    return render(request, 'generate.html')

def excel(request):
    SAVE_NAME = str(random.randint(1, 1001)) + '.xlsx'
    DIR = os.path.abspath(str(BASE_DIR) + '/static/temp_files/' + SAVE_NAME)

    xmlfile = xlsxwriter.Workbook(DIR)
    sheet = xmlfile.add_worksheet('Data')

    sheet.write(0, 0,  'id')
    sheet.write(0, 1, 'Формат')
    sheet.write(0, 2, 'Номинация')
    sheet.write(0, 3, 'Участник')
    sheet.write(0, 4, 'Преподаватель')
    sheet.write(0, 5, 'Класс')
    sheet.write(0, 6, 'Концертмейстер')
    sheet.write(0, 7, 'Заведение')
    sheet.write(0, 8, 'Программа')
    sheet.write(0, 9, 'Ссылка')
    sheet.write(0, 10, 'Место')

    claims = Claim.objects.all()

    for row, claim in enumerate(claims):
        sheet.write(row + 1, 0, claim.id)
        sheet.write(row + 1, 1, claim.type_claim)
        sheet.write(row + 1, 2, claim.select_claim)
        sheet.write(row + 1, 3, claim.name_claim)
        sheet.write(row + 1, 4, claim.parent_claim)
        sheet.write(row + 1, 5, claim.class_claim)
        sheet.write(row + 1, 6, claim.concert_claim)
        sheet.write(row + 1, 7, claim.city_claim)
        sheet.write(row + 1, 8, claim.prog_claim)
        sheet.write(row + 1, 9, claim.url_claim)
        sheet.write(row + 1, 10, '')

    xmlfile.close()
    return render(request, 'excel.html', context={'dir': SAVE_NAME})

def word(request):
    claims = Claim.objects.all()
    dirs = create_report(claims)
    return render(request, 'word.html', context={'dir': dirs})

def apply(request):
    if request.method == 'POST':
        type_cl = request.POST.get('type_claim')
        select_cl = request.POST.get('select_claim')
        name_cl = request.POST.get('name_claim')
        parent_cl = request.POST.get('parent_claim')
        class_cl = request.POST.get('class_claim')
        concert_cl = request.POST.get('concert_claim')
        city_cl = request.POST.get('city_claim')
        prog_cl = request.POST.get('prog_claim')
        url_cl = request.POST.get('url_claim')

        data = {'type': type_cl, 'select': select_cl, 'name': name_cl,
                'parent': parent_cl, 'class': class_cl, 'concert': concert_cl,
                'city': city_cl, 'prog': prog_cl, 'url': url_cl
                }
        claim = Claim.objects.create(type_claim=type_cl,
                                    select_claim=select_cl,
                                    name_claim=name_cl,
                                    parent_claim=parent_cl,
                                    class_claim=class_cl,
                                    concert_claim=concert_cl,
                                    city_claim=city_cl,
                                    prog_claim=prog_cl,
                                    url_claim=url_cl,
                                    place_claim=0)
        data['id'] = claim.id
        return render(request, 'apply.html', context=data)
    else:
        return HttpResponse('Error')